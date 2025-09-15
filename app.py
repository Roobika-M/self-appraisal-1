from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import os
import json
import pandas as pd
from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import threading
from datetime import datetime

app = Flask(__name__, static_folder="build", static_url_path="/")
CORS(app, origins=['http://localhost:8080'], supports_credentials=True)
app.secret_key = "your_secret_key"

# Globals
staffname = ""
detaillist = []
excel_path = ""
research = selfm = mentor = academics = hod = 0

@app.route("/")
def home():
    # Serve React index.html for root
    return send_from_directory(app.static_folder, "index.html")

@app.route("/login", methods=["POST", "GET"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        if not username or not password:
            return jsonify({"success": False, "error": "Please enter both username and password."}), 400
        # dummy auth: password 'admin' allows access
        if password == "admin":
            return jsonify({"success": True}), 200
        else:
            return jsonify({"success": False, "error": "Please enter correct username and password."}), 401
    return jsonify({"message": "Login endpoint"})

@app.route("/upload", methods=["POST"])
def upload():
    global staffname, detaillist, excel_path
    name = request.form.get("name")
    designation = request.form.get("designation")
    department = request.form.get("department")
    emp_id = request.form.get("employee_id")

    # Log received form data
    print("Received form data:", {
        "name": name,
        "designation": designation,
        "department": department,
        "employee_id": emp_id
    })

    if not all([name, designation, department, emp_id]):
        print("Missing required form fields.")
        return jsonify({"success": False, "error": "Please fill in all details."}), 400
    staffname = name
    detaillist = [name, designation, department, emp_id]

    excel_file = request.files.get("excel_file")
    template_file = request.files.get("word_file")

    # Log received files
    print("Received files:", {
        "excel_file": excel_file.filename if excel_file else None,
        "word_file": template_file.filename if template_file else None
    })

    if not excel_file:
        print("Missing Excel file.")
        return jsonify({"error": "Excel file is required."}), 400

    upload_folder = os.getcwd()
    excel_path = os.path.join(upload_folder, excel_file.filename)
    # Always use template.docx from project folder
    template_path = os.path.join(upload_folder, "template.docx")

    try:
        excel_file.save(excel_path)
    except Exception as e:
        print(f"Error saving Excel file: {e}")
        return jsonify({"success": False, "error": f"File save failed: {str(e)}"}), 500

    try:
        processing(excel_path, staffname, template_path)
    except Exception as e:
        print(f"Error in processing: {e}")
        return jsonify({"success": False, "error": f"Processing failed: {str(e)}"}), 500

    # Save appraisal to history
    with history_lock:
        history = load_history()
        try:
            total_score = int(research) + int(selfm) + int(mentor) + int(academics) + int(hod)
        except Exception:
            total_score = 0
        appraisal = {
            "name": staffname,
            "designation": detaillist[1] if len(detaillist) > 1 else "",
            "dept": detaillist[2] if len(detaillist) > 2 else "",
            "empid": detaillist[3] if len(detaillist) > 3 else "",
            "research": research,
            "selfm": selfm,
            "mentor": mentor,
            "academics": academics,
            "hod": hod,
            "total_score": total_score,
            "timestamp": datetime.now().isoformat()
        }
        history.append(appraisal)
        save_history(history)

    print("File processed successfully.")
    return jsonify({"success": True, "message": "File processed successfully."}), 200


@app.route("/download/<file_type>", methods=["GET"])
def download(file_type):
    base = os.getcwd()
    if file_type == "docx":
        file_path = os.path.join(base, "filled_template.docx")
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
        return send_file(file_path, as_attachment=True)
    elif file_type == "pdf":
        docx_path = os.path.join(base, "filled_template.docx")
        pdf_path = os.path.join(base, "filled_template.pdf")
        if not os.path.exists(docx_path):
            return jsonify({"error": "DOCX file not found for conversion"}), 404
        # Convert DOCX to PDF if PDF doesn't exist or is older than DOCX
        try:
            if not os.path.exists(pdf_path) or os.path.getmtime(pdf_path) < os.path.getmtime(docx_path):
                from docx2pdf import convert
                convert(docx_path, pdf_path)
        except Exception as e:
            print(f"Error converting DOCX to PDF: {e}")
            return jsonify({"error": f"PDF conversion failed: {str(e)}"}), 500
        if not os.path.exists(pdf_path):
            return jsonify({"error": "PDF file not found after conversion"}), 404
        return send_file(pdf_path, as_attachment=True)
    return jsonify({"error": "Invalid file type"}), 400


HISTORY_FILE = "appraisal_history.json"
history_lock = threading.Lock()

def load_history():
    if not os.path.exists(HISTORY_FILE):
        return []
    with open(HISTORY_FILE, "r", encoding="utf-8") as f:
        try:
            return json.load(f)
        except Exception:
            return []

def save_history(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)

@app.route("/download_path")
def download_path():
    history = load_history()
    return jsonify(history)

def find_header_row(excel_path, sheet_name):
    """Attempt to find the header row index where 'Faculty Name' or similar exists.
       Returns None if not found."""
    try:
        df_temp = pd.read_excel(excel_path, sheet_name=sheet_name, nrows=15)
    except Exception as e:
        print(f"Error reading sheet {sheet_name} to detect header: {e}")
        return None

    for idx, row in df_temp.iterrows():
        row_values = [str(val).lower().strip() for val in row if pd.notna(val)]
        if "faculty name" in row_values or "name of the faculty" in row_values:
            return idx
    return None


def processing(excel_path, staffname, template_path):
    """Full processing: read provided Excel, populate template Word docs and compute scores."""
    global research, selfm, mentor, academics, hod, detaillist

    # Initialize detail list if missing
    if not detaillist:
        detaillist = [staffname, "", "", ""]

    # Load template and corrective-action doc
    try:
        doc = Document(template_path)
        doc2=Document()
    except Exception as e:
        raise RuntimeError(f"Could not open template docx at {template_path}: {e}")

    corrective_doc_path = "Faculty Appraisal- Corrective Action Report.docx"
    try:
        fdoc = Document(corrective_doc_path)
    except Exception as e:
        # If corrective doc not found, create a copy of template for corrective operations
        print(f"Corrective doc not found or couldn't open: {e}. Using template as fallback.")
        fdoc = Document(template_path)

    # reset scores and counters
    research = 0
    selfm = 0
    mentor = 0
    academics = 0
    hod = 0

    # initialize r and p and s counters used across different sections
    r_counters = {f"r{i}_1": 0 for i in range(1, 14)}
    p_counters = {f"p{i}_1": 0 for i in range(1, 8)}
    s_counters = {f"s{i}_1": 0 for i in range(1, 6)}

    # Helper functions
    def safe_float(x):
        try:
            return float(x)
        except Exception:
            return 0.0

    def get_grade_1(value):
        if value > 95:
            return 5
        elif 90 <= value <= 95:
            return 4
        elif 80 <= value < 90:
            return 3
        elif 70 <= value < 80:
            return 2
        elif 60 <= value < 70:
            return 1
        elif 50 <= value < 60:
            return 0
        else:
            return -1

    def get_grade_2(value):
        if 0 < value <= 2:
            return 1
        elif 3 <= value <= 4:
            return 2
        elif 5 <= value <= 6:
            return 3
        elif 7 <= value <= 9:
            return 4
        else:
            return 5

    def get_grade_negative(value):
        if 0 < value <= 10:
            return -1
        elif 11 <= value <= 20:
            return -2
        elif 21 <= value <= 30:
            return -3
        elif 31 <= value <= 40:
            return -4
        else:
            return -5

    # read list of sheets safely
    try:
        sheet_names = pd.ExcelFile(excel_path).sheet_names
    except Exception as e:
        sheet_names = []
        print("Could not read Excel file or sheets:", e)

    ############### Academics section (copy table and compute totals) ###############
    try:
        # We assume doc.tables[1] is destination and uploaded doc1's table[1] is source in original logic.
        # Since user uploaded only one template, attempt to use template's own tables as both source/dest.
        source_table = doc2.tables[1]
        destination_table = doc.tables[1]
    except Exception:
        source_table = None
        destination_table = None

    if source_table is not None and destination_table is not None:
        scores = [0.0] * 6
        nos = 0
        # iterate rows starting from 2 as in original code
        for i in range(2, len(source_table.rows)):
            row = source_table.rows[i]
            first_cell = row.cells[0].text.strip()

            if first_cell.lower() == "total/average":
                # compute nos as number of valid data rows
                nos = max(0, i - 3)
                # Add two rows for Total/Average and Marks
                for j in range(i, i + 2):
                    if j >= len(destination_table.rows):
                        destination_table.add_row()
                    new_row = destination_table.rows[j]
                    source_row = source_table.rows[j] if j < len(source_table.rows) else None

                    # ensure correct number of cells
                    while len(new_row.cells) < (len(source_row.cells) if source_row else 0):
                        new_row._tr.add_tc()

                    # merge first 4 cells into one label cell (if possible)
                    try:
                        merged_cell = new_row.cells[0].merge(new_row.cells[1])
                        merged_cell = merged_cell.merge(new_row.cells[2])
                        merged_cell = merged_cell.merge(new_row.cells[3])
                    except Exception:
                        pass

                    if j == i:
                        # Total/Average row
                        try:
                            new_row.cells[3].text = "Total/Average"
                            avg = (scores[0] / nos) if nos > 0 else 0
                            # put average at col 4 and remaining sums next
                            if len(new_row.cells) > 4:
                                new_row.cells[4].text = f"{avg:.2f}"
                            for k in range(1, 6):
                                if 4 + k < len(new_row.cells):
                                    new_row.cells[4 + k].text = f"{scores[k]:.2f}"
                        except Exception:
                            pass
                    else:
                        # Marks row based on grading functions
                        try:
                            new_row.cells[3].text = "Marks(Ref guideline for awarding score)"
                            if nos > 0:
                                val0 = scores[0] / nos
                            else:
                                val0 = 0
                            if len(new_row.cells) > 4:
                                new_row.cells[4].text = str(get_grade_1(val0))
                            if len(new_row.cells) > 5:
                                new_row.cells[5].text = str(get_grade_2(scores[1]))
                            if len(new_row.cells) > 6:
                                new_row.cells[6].text = str(get_grade_2(scores[2]))
                            if len(new_row.cells) > 7:
                                new_row.cells[7].text = str(get_grade_2(scores[3]))
                            if len(new_row.cells) > 8:
                                new_row.cells[8].text = str(get_grade_negative(scores[4]))
                            if len(new_row.cells) > 9:
                                new_row.cells[9].text = str(get_grade_negative(scores[5]))
                        except Exception:
                            pass
                # stop scanning source table
                break

            # Normal copy + accumulate for numeric columns 4..9
            if i >= len(destination_table.rows):
                destination_table.add_row()
            new_row = destination_table.rows[i]
            # ensure enough cells
            while len(new_row.cells) < len(row.cells):
                new_row._tr.add_tc()
            for j in range(len(row.cells)):
                text = row.cells[j].text.strip()
                # copy text
                try:
                    new_row.cells[j].text = text
                except Exception:
                    pass
                # accumulate numeric columns (4..9)
                if 4 <= j <= 9:
                    scores[j - 4] += safe_float(text)

        # academics score computed via helper get_total_academics_score
        def get_total_academics_score(scores_list, nos_val):
            total = 0
            if nos_val > 0:
                total += get_grade_1(scores_list[0] / nos_val)
            total += get_grade_2(scores_list[1])
            total += get_grade_2(scores_list[2])
            total += get_grade_2(scores_list[3])
            total += get_grade_negative(scores_list[4])
            total += get_grade_negative(scores_list[5])
            return total

        academics = get_total_academics_score(scores, nos)
    else:
        academics = 0

    ############### Publications: Journals, Books ###############
    # Initialize many counters used in original code
    for k in range(1, 14):
        r_counters[f"r{k}_1"] = 0
    # Also ensure r9 used earlier gets a safe var if present
    r_counters["r9"] = r_counters.get("r9", 0)

    # Journals
    if "Journal Publication" in sheet_names:
        header = find_header_row(excel_path, "Journal Publication")
        skiprows = header + 1 if header is not None else 0
        try:
            df_journal = pd.read_excel(excel_path, sheet_name="Journal Publication", skiprows=skiprows)
            df_journal.columns = df_journal.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selected_col = next((col for col in possible_names if col in df_journal.columns), None)
            if selected_col:
                df_journal[selected_col] = df_journal[selected_col].ffill()
                df_filtered = df_journal[df_journal[selected_col].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Journal Publication:", e)

        if not df_filtered.empty:
            # index for journal table in template (preserved from original)
            table3_index = 3
            if table3_index < len(doc.tables):
                table3 = doc.tables[table3_index]
                start_row = 1
                n = 0
                m_val = 0
                for i, (_, row) in enumerate(df_filtered.iterrows()):
                    row_index = start_row + i
                    if row_index + 1 >= len(table3.rows):
                        table3.add_row()
                    # write to row_index+1 as original did
                    target_row = row_index + 1
                    try:
                        table3.cell(target_row, 0).text = str(i + 1)
                        table3.cell(target_row, 1).text = str(row.get("Paper Title", "-"))
                        table3.cell(target_row, 2).text = str(row.get("Journal Name", "-"))
                        table3.cell(target_row, 3).text = str(row.get("Year of Publication", "-"))
                        table3.cell(target_row, 4).text = str(row.get("ISSN", "-"))
                        table3.cell(target_row, 5).text = str(row.get("Web Link", "-"))
                        impact = row.get("Impact Factor", "-")
                        table3.cell(target_row, 6).text = str(impact)
                        # scoring rules
                        try:
                            if pd.notna(impact) and impact != "-":
                                if float(impact) > 3:
                                    n += 3
                                    r_counters["r2_1"] += 3
                                elif 1.5 < float(impact) <= 3:
                                    n += 2
                                    r_counters["r3_1"] += 2
                                elif 1 <= float(impact) <= 1.5:
                                    n += 1
                                    r_counters["r4_1"] += 1
                        except Exception:
                            pass
                        # base points per journal entry
                        n += 2
                        m_val += 2
                    except Exception as e:
                        print("Error filling journal row:", e)
                # Add total row
                try:
                    table3.add_row()
                    last = table3.rows[-1]
                    # merge first to second-last cells for label
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0] if len(last.cells) >= 2 else last.cells[0].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    # put total score in last cell
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
                research += n

    # Books
    if "Book Publication" in sheet_names:
        header = find_header_row(excel_path, "Book Publication")
        skiprows = header + 1 if header is not None else 0
        try:
            df_bookpub = pd.read_excel(excel_path, sheet_name="Book Publication", skiprows=skiprows)
            df_bookpub.columns = df_bookpub.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selected_col = next((col for col in possible_names if col in df_bookpub.columns), None)
            if selected_col:
                df_bookpub[selected_col] = df_bookpub[selected_col].ffill()
                df_filtered = df_bookpub[df_bookpub[selected_col].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Book Publication:", e)

        if not df_filtered.empty:
            table4_index = 4
            if table4_index < len(doc.tables):
                table4 = doc.tables[table4_index]
                start_row = 1
                n = 0
                for i, (_, row) in enumerate(df_filtered.iterrows()):
                    row_index = start_row + i
                    if row_index + 1 >= len(table4.rows):
                        table4.add_row()
                    try:
                        table4.cell(row_index + 1, 0).text = str(i + 1)
                        table4.cell(row_index + 1, 1).text = str(row.get("Book Title", "-"))
                        table4.cell(row_index + 1, 2).text = str(row.get("Publication Name", "-"))
                        table4.cell(row_index + 1, 3).text = str(row.get("Date of Publication", "-"))
                        table4.cell(row_index + 1, 4).text = str(row.get("ISBN", "-"))
                        table4.cell(row_index + 1, 5).text = str(row.get("Description", "-"))
                        # basic scoring for books if desired
                        n += 1
                    except Exception as e:
                        print("Error filling book row:", e)
                try:
                    table4.add_row()
                    last = table4.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
                research += n

    ############### Conferences ###############
    if "Conferences" in sheet_names:
        header = find_header_row(excel_path, "Conferences")
        skiprows = header + 1 if header is not None else 0
        try:
            df_conference = pd.read_excel(excel_path, sheet_name="Conferences", skiprows=skiprows)
            df_conference.columns = df_conference.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selected_col = next((col for col in possible_names if col in df_conference.columns), None)
            if selected_col:
                df_conference[selected_col] = df_conference[selected_col].ffill()
                df_filtered = df_conference[df_conference[selected_col].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Conferences:", e)

        if not df_filtered.empty:
            table6_index = 6
            table7_index = 7
            if table6_index < len(doc.tables):
                table6 = doc.tables[table6_index]
            else:
                table6 = None
            if table7_index < len(doc.tables):
                table7 = doc.tables[table7_index]
            else:
                table7 = None

            start_row = 1
            n_total = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                conference_type = str(row.get("Conference Type", "")).strip().lower()
                if conference_type == "international":
                    table = table6
                    n = 2
                    r_counters["r8_1"] = r_counters.get("r8_1", 0) + 2
                else:
                    table = table7
                    n = 1
                    r_counters["r9_1"] = r_counters.get("r9_1", 0) + 1
                if table is None:
                    continue
                row_index = start_row + i
                if row_index + 1 >= len(table.rows):
                    table.add_row()
                try:
                    table.cell(row_index + 1, 0).text = str(i + 1)
                    table.cell(row_index + 1, 1).text = str(row.get("Paper Title", "-"))
                    table.cell(row_index + 1, 2).text = str(row.get("Organized By", "-"))
                    table.cell(row_index + 1, 3).text = str(row.get("From Date", "-"))
                    table.cell(row_index + 1, 4).text = str(row.get("Place", "-"))
                    table.cell(row_index + 1, 5).text = str(row.get("Role", "-"))
                except Exception as e:
                    print("Error filling conference row:", e)
                n_total += n
            # add totals to whichever table was used last
            if table6 is not None or table7 is not None:
                tbl = table6 if table6 is not None else table7
                try:
                    tbl.add_row()
                    last = tbl.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n_total)
                except Exception:
                    pass
                research += n_total

    ############### Research Grants & Seminars (Research Grant sheet reused) ###############
    if "Research Grant" in sheet_names:
        header = find_header_row(excel_path, "Research Grant")
        skiprows = header + 1 if header is not None else 0
        try:
            df_research = pd.read_excel(excel_path, sheet_name="Research Grant", skiprows=skiprows)
            df_research.columns = df_research.columns.str.strip()
            # ensure column names tolerant
            if "Faculty Name" in df_research.columns:
                df_research["Faculty Name"] = df_research["Faculty Name"].ffill()
                df_filtered = df_research[df_research["Faculty Name"].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Research Grant:", e)

        # Grants where Coordinator == 'applied' => treated as applications (table7 in original)
        if not df_filtered.empty:
            table7_index = 7
            if table7_index < len(doc.tables):
                table7 = doc.tables[table7_index]
            else:
                table7 = None

            n = 0
            total_amt = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                row_index = 1 + i
                coord = str(row.get("Coordinator", "-")).strip().lower()
                amount = row.get("Amount", 0)
                try:
                    if coord == "applied":
                        if table7 is not None:
                            if row_index + 1 >= len(table7.rows):
                                table7.add_row()
                            table7.cell(row_index + 1, 0).text = str(i + 1)
                            table7.cell(row_index + 1, 1).text = str(row.get("Coordinator", "-"))
                            table7.cell(row_index + 1, 2).text = str(row.get("Title", "-"))
                            table7.cell(row_index + 1, 3).text = str(row.get("Type", "-"))
                            table7.cell(row_index + 1, 4).text = str(row.get("Funding Agent", "-"))
                            table7.cell(row_index + 1, 5).text = str(amount)
                            table7.cell(row_index + 1, 6).text = str(row.get("Applied On", "-"))
                        if pd.notna(amount) and amount != "-" and safe_float(amount) > 0:
                            total_amt += safe_float(amount)
                except Exception as e:
                    print("Error filling research grant row:", e)

            # scoring for grant amounts (example from original: > 1,000,000 gives points)
            try:
                if total_amt > 1000000:
                    add_pts = int(total_amt // 1000000) * 2
                    n += add_pts
                    r_counters["r10_1"] = r_counters.get("r10_1", 0) + add_pts
            except Exception:
                pass

            # add totals row
            if table7 is not None:
                try:
                    table7.add_row()
                    last = table7.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            research += n

        # Seminars (same sheet but coordinator != applied) -> table9 in original
        if not df_filtered.empty:
            table9_index = 9
            if table9_index < len(doc.tables):
                table9 = doc.tables[table9_index]
            else:
                table9 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                coord = str(row.get("Coordinator", "-")).strip().lower()
                if coord != "applied":
                    row_index = 1 + i
                    if table9 is not None:
                        if row_index + 1 >= len(table9.rows):
                            table9.add_row()
                        try:
                            table9.cell(row_index + 1, 0).text = str(i + 1)
                            table9.cell(row_index + 1, 1).text = str(row.get("Coordinator", "-"))
                            table9.cell(row_index + 1, 2).text = str(row.get("Title", "-"))
                            table9.cell(row_index + 1, 3).text = str(row.get("Type", "-"))
                            table9.cell(row_index + 1, 4).text = str(row.get("Funding Agent", "-"))
                            table9.cell(row_index + 1, 5).text = str(row.get("Amount", "-"))
                            table9.cell(row_index + 1, 6).text = str(row.get("Applied On", "-"))
                            amt = safe_float(row.get("Amount", 0))
                            if amt > 50000:
                                inc = int(amt // 50000)
                                n += inc
                                r_counters["r11_1"] = r_counters.get("r11_1", 0) + inc
                        except Exception as e:
                            print("Error filling seminar row:", e)
            if table9 is not None:
                try:
                    table9.add_row()
                    last = table9.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            research += n

    ############### Patents ###############
    if "Patents" in sheet_names:
        try:
            df_patent = pd.read_excel(excel_path, sheet_name="Patents")
            df_patent.columns = df_patent.columns.str.strip()
            # tolerant column name
            possible_name_cols = ["Faculty name", "Faculty Name", "Faculty"]
            selected_col = next((c for c in possible_name_cols if c in df_patent.columns), None)
            if selected_col:
                df_patent[selected_col] = df_patent[selected_col].ffill()
                df_filtered = df_patent[df_patent[selected_col].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Patents:", e)

        if not df_filtered.empty:
            table10_index = 10
            if table10_index < len(doc.tables):
                table10 = doc.tables[table10_index]
            else:
                table10 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                row_index = 1 + i
                if table10 is not None and row_index + 1 >= len(table10.rows):
                    table10.add_row()
                try:
                    # Serial
                    if table10 is not None:
                        table10.cell(row_index + 1, 0).text = str(i + 1)
                        title = str(row.get("Title", "-"))
                        table10.cell(row_index + 1, 1).text = title
                        status = str(row.get("Status", "")).strip().lower()
                        date_value = str(row.get("Date", "-"))
                        if status == "filed":
                            table10.cell(row_index + 1, 2).text = date_value  # filing
                            table10.cell(row_index + 1, 3).text = "-"
                        elif status == "published":
                            table10.cell(row_index + 1, 2).text = "-"  # filing
                            table10.cell(row_index + 1, 3).text = date_value  # published
                            n += 5
                            r_counters["r12_1"] = r_counters.get("r12_1", 0) + 5
                        else:
                            table10.cell(row_index + 1, 2).text = "-"
                            table10.cell(row_index + 1, 3).text = "-"
                        table10.cell(row_index + 1, 4).text = str(row.get("Status", "-"))
                except Exception as e:
                    print("Error filling patent row:", e)
            if table10 is not None:
                try:
                    table10.add_row()
                    last = table10.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            research += n

    ############### Workshops, Internships, MOOC, MoU, etc. ###############
    # Workshop (table13_index in original was 14)
    if "Workshop" in sheet_names:
        header = find_header_row(excel_path, "Workshop")
        skiprows = header + 1 if header is not None else 0
        try:
            df_workshop = pd.read_excel(excel_path, sheet_name="Workshop", skiprows=skiprows)
            df_workshop.columns = df_workshop.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selected_col = next((col for col in possible_names if col in df_workshop.columns), None)
            if selected_col:
                df_workshop[selected_col] = df_workshop[selected_col].ffill()
                df_filtered = df_workshop[df_workshop[selected_col].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Workshop:", e)

        if not df_filtered.empty:
            table13_index = 14
            if table13_index < len(doc.tables):
                table13 = doc.tables[table13_index]
            else:
                table13 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                role = str(row.get("Role", "")).strip().lower()
                if role == "attended":
                    if table13 is not None and (1 + i + 1) >= len(table13.rows):
                        table13.add_row()
                    try:
                        table13.cell(1 + i + 1, 0).text = str(i + 1)
                        table13.cell(1 + i + 1, 1).text = str(row.get("Topic", "-"))
                        from_date = str(row.get("From Date", "-"))
                        to_date = str(row.get("To Date", "-"))
                        table13.cell(1 + i + 1, 2).text = f"{from_date} to {to_date}"
                        table13.cell(1 + i + 1, 3).text = str(row.get("Description", "-"))
                        table13.cell(1 + i + 1, 4).text = str(row.get("Venue", "-"))
                    except Exception:
                        pass
                    if n < 3:
                        n += 1
                        p_counters["p1_1"] = p_counters.get("p1_1", 0) + 1
            if table13 is not None:
                try:
                    table13.add_row()
                    last = table13.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            selfm += n

    # Faculty Internship (table14)
    if "Faculty Internship" in sheet_names:
        header = find_header_row(excel_path, "Faculty Internship")
        skiprows = header + 1 if header is not None else 0
        try:
            df_develop = pd.read_excel(excel_path, sheet_name="Faculty Internship", skiprows=skiprows)
            df_develop.columns = df_develop.columns.str.strip()
            # possible column names
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selectedcol = next((col for col in possible_names if col in df_develop.columns), None)
            if selectedcol:
                df_develop[selectedcol] = df_develop[selectedcol].ffill()
                df_filtered = df_develop[df_develop[selectedcol].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Faculty Internship:", e)

        if not df_filtered.empty:
            table14_index = 15
            if table14_index < len(doc.tables):
                table14 = doc.tables[table14_index]
            else:
                table14 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                if table14 is not None and (1 + i + 1) >= len(table14.rows):
                    table14.add_row()
                try:
                    table14.cell(1 + i + 1, 0).text = str(i + 1)
                    table14.cell(1 + i + 1, 1).text = str(row.get("FDP Name", "-"))
                    from_date = str(row.get("From Date", "-"))
                    to_date = str(row.get("To Date", "-"))
                    table14.cell(1 + i + 1, 2).text = f"{from_date} to {to_date}"
                    table14.cell(1 + i + 1, 3).text = str(row.get("Description", "-"))
                    table14.cell(1 + i + 1, 4).text = str(row.get("National or International", "-"))
                    n += 3
                    p_counters["p2_1"] = p_counters.get("p2_1", 0) + 3
                except Exception:
                    pass
            if table14 is not None:
                try:
                    table14.add_row()
                    last = table14.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            selfm += n

    # MOOC Course (table15)
    if "MOOC Course" in sheet_names:
        header = find_header_row(excel_path, "MOOC Course")
        skiprows = header + 1 if header is not None else 0
        try:
            df_mooc = pd.read_excel(excel_path, sheet_name="MOOC Course", skiprows=skiprows)
            df_mooc.columns = df_mooc.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selectedcol = next((col for col in possible_names if col in df_mooc.columns), None)
            if selectedcol:
                df_mooc[selectedcol] = df_mooc[selectedcol].ffill()
                df_filtered = df_mooc[df_mooc[selectedcol].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading MOOC Course:", e)

        if not df_filtered.empty:
            table15_index = 16
            if table15_index < len(doc.tables):
                table15 = doc.tables[table15_index]
            else:
                table15 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                if table15 is not None and (1 + i + 1) >= len(table15.rows):
                    table15.add_row()
                try:
                    table15.cell(1 + i + 1, 0).text = str(i + 1)
                    table15.cell(1 + i + 1, 1).text = str(row.get("Coure Title", "-"))
                    table15.cell(1 + i + 1, 2).text = str(row.get("Course Type", "-"))
                    from_date = str(row.get("From Date", "-"))
                    to_date = str(row.get("To Date", "-"))
                    table15.cell(1 + i + 1, 3).text = f"{from_date} to {to_date}"
                    table15.cell(1 + i + 1, 4).text = str(row.get("Duration", "-"))
                    table15.cell(1 + i + 1, 5).text = str(row.get("Awards", "-"))
                    if n < 4:
                        n += 2
                        p_counters["p3_1"] = p_counters.get("p3_1", 0) + 2
                except Exception:
                    pass
            if table15 is not None:
                try:
                    table15.add_row()
                    last = table15.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            selfm += n

    # MoU (table16)
    if "MoU" in sheet_names:
        header = find_header_row(excel_path, "MoU")
        skiprows = header + 1 if header is not None else 0
        try:
            df_mou = pd.read_excel(excel_path, sheet_name="MoU", skiprows=skiprows)
            df_mou.columns = df_mou.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selectedcol = next((col for col in possible_names if col in df_mou.columns), None)
            if selectedcol:
                df_mou[selectedcol] = df_mou[selectedcol].ffill()
                df_filtered = df_mou[df_mou[selectedcol].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading MoU:", e)

        if not df_filtered.empty:
            table16_index = 17
            if table16_index < len(doc.tables):
                table16 = doc.tables[table16_index]
            else:
                table16 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                if table16 is not None and (1 + i + 1) >= len(table16.rows):
                    table16.add_row()
                try:
                    table16.cell(1 + i + 1, 0).text = str(i + 1)
                    table16.cell(1 + i + 1, 1).text = str(row.get(selectedcol, "-"))
                    table16.cell(1 + i + 1, 2).text = str(row.get("Company Name", "-"))
                    from_date = str(row.get("From Date", "-"))
                    to_date = str(row.get("To Date", "-"))
                    table16.cell(1 + i + 1, 3).text = f"{from_date} to {to_date}"
                    table16.cell(1 + i + 1, 4).text = str(row.get("Industry SPOC", "-"))
                    table16.cell(1 + i + 1, 5).text = str(row.get("Duration", "-"))
                    n += 1
                    p_counters["p4_1"] = p_counters.get("p4_1", 0) + 1
                except Exception:
                    pass
            if table16 is not None:
                try:
                    table16.add_row()
                    last = table16.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            selfm += n

    ############### Workshops conducted (Workshops sheet) ###############
    if "Workshops" in sheet_names:
        header = find_header_row(excel_path, "Workshops")
        skiprows = header + 1 if header is not None else 0
        try:
            df_workshops_conducted = pd.read_excel(excel_path, sheet_name="Workshops", skiprows=skiprows)
            df_workshops_conducted.columns = df_workshops_conducted.columns.str.strip()
            possible_names = ["Faculty Name", "Faculty name", "Name of the Faculty", "Name", "Faculty"]
            selectedcol = next((col for col in possible_names if col in df_workshops_conducted.columns), None)
            if selectedcol:
                df_workshops_conducted[selectedcol] = df_workshops_conducted[selectedcol].ffill()
                df_filtered = df_workshops_conducted[
                    (df_workshops_conducted[selectedcol].astype(str).str.strip() == staffname) &
                    (df_workshops_conducted["Role"].fillna("").astype(str).str.strip().str.lower() == "conducted")
                ]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Workshops:", e)

        if not df_filtered.empty:
            table17_index = 19
            if table17_index < len(doc.tables):
                table17 = doc.tables[table17_index]
            else:
                table17 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                if table17 is not None and (1 + i + 1) >= len(table17.rows):
                    table17.add_row()
                try:
                    table17.cell(1 + i + 1, 0).text = str(i + 1)
                    table17.cell(1 + i + 1, 1).text = str(row.get("Topic", "-"))
                    table17.cell(1 + i + 1, 2).text = str(row.get("Department", "-"))
                    from_date = str(row.get("From Date", "-"))
                    to_date = str(row.get("To Date", "-"))
                    table17.cell(1 + i + 1, 3).text = f"{from_date} to {to_date}"
                    table17.cell(1 + i + 1, 4).text = str(row.get("No of Students", "-"))
                    table17.cell(1 + i + 1, 5).text = str(row.get("Venue", "-"))
                    table17.cell(1 + i + 1, 6).text = str(row.get("Description", "-"))
                    n += 0.5
                    p_counters["p6_1"] = p_counters.get("p6_1", 0) + 0.5
                except Exception:
                    pass
            if table17 is not None:
                try:
                    table17.add_row()
                    last = table17.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            selfm += n

    ############### Guest Lectures / Expert Visits ###############
    if "Guest Lectures" in sheet_names:
        header = find_header_row(excel_path, "Guest Lectures")
        skiprows = header + 1 if header is not None else 0
        try:
            df_experts = pd.read_excel(excel_path, sheet_name="Guest Lectures", skiprows=skiprows)
            df_experts.columns = df_experts.columns.str.strip()
            if "Faculty Name" in df_experts.columns:
                df_experts["Faculty Name"] = df_experts["Faculty Name"].ffill()
                df_filtered = df_experts[df_experts["Faculty Name"].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Guest Lectures:", e)

        if not df_filtered.empty:
            table19_index = 20
            if table19_index < len(doc.tables):
                table19 = doc.tables[table19_index]
            else:
                table19 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                if table19 is not None and (1 + i + 1) >= len(table19.rows):
                    table19.add_row()
                try:
                    table19.cell(1 + i + 1, 0).text = str(i + 1)
                    table19.cell(1 + i + 1, 1).text = str(row.get("Chief Guest Name", "-"))
                    table19.cell(1 + i + 1, 2).text = str(row.get("Address", "-"))
                    table19.cell(1 + i + 1, 3).text = str(row.get("Topic Name", "-"))
                    from_date = str(row.get("From Date", "-"))
                    to_date = str(row.get("To Date", "-"))
                    table19.cell(1 + i + 1, 4).text = f"{from_date} to {to_date}"
                    table19.cell(1 + i + 1, 5).text = str(row.get("Description", "-"))
                    table19.cell(1 + i + 1, 6).text = str(row.get("Topic Delivered", "-"))
                    n += 1
                    p_counters["p7_1"] = p_counters.get("p7_1", 0) + 1
                except Exception:
                    pass
            if table19 is not None:
                try:
                    table19.add_row()
                    last = table19.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            selfm += n

    ############### Projects Guided / Mentoring ###############
    if "Project Guided or Mentoring" in sheet_names:
        header = find_header_row(excel_path, "Project Guided or Mentoring")
        skiprows = header + 1 if header is not None else 0
        try:
            df_project = pd.read_excel(excel_path, sheet_name="Project Guided or Mentoring", skiprows=skiprows)
            df_project.columns = df_project.columns.str.strip()
            if "Faculty Name" in df_project.columns:
                df_project["Faculty Name"] = df_project["Faculty Name"].ffill()
                df_filtered = df_project[df_project["Faculty Name"].astype(str).str.strip() == staffname]
            else:
                df_filtered = pd.DataFrame()
        except Exception as e:
            df_filtered = pd.DataFrame()
            print("Error reading Project Guided or Mentoring:", e)

        if not df_filtered.empty:
            table21_index = 22
            if table21_index < len(doc.tables):
                table21 = doc.tables[table21_index]
            else:
                table21 = None
            n = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                if table21 is not None and (1 + i + 1) >= len(table21.rows):
                    table21.add_row()
                try:
                    table21.cell(1 + i + 1, 0).text = str(i + 1)
                    table21.cell(1 + i + 1, 1).text = str(row.get("Project Title", "-"))
                    table21.cell(1 + i + 1, 2).text = str(row.get("Number of Students", "-"))
                    table21.cell(1 + i + 1, 3).text = str(row.get("Title of Hackathon", "-"))
                    table21.cell(1 + i + 1, 4).text = str(row.get("Organized By", "-"))
                    table21.cell(1 + i + 1, 5).text = str(row.get("Date", "-"))
                    table21.cell(1 + i + 1, 6).text = str(row.get("Status", "-"))
                    # example scoring
                    n = 1
                    s_counters["s1_1"] = 1
                except Exception:
                    pass
            if table21 is not None:
                try:
                    table21.add_row()
                    last = table21.rows[-1]
                    try:
                        last.cells[0].merge(last.cells[-2])
                    except Exception:
                        pass
                    paragraph = last.cells[-2].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    last.cells[-1].text = str(n)
                except Exception:
                    pass
            mentor += n

    # Prepare placeholders for main template and corrective doc
    placeholders = {
        "{{research}}": str(research),
        "{{self}}": str(selfm),
        "{{mentorship}}": str(mentor),
        "{{academics}}": str(academics),
        "{{name}}": detaillist[0] if detaillist and len(detaillist) > 0 else staffname,
        "{{designation}}": detaillist[1] if len(detaillist) > 1 else "",
        "{{dept}}": detaillist[2] if len(detaillist) > 2 else "",
        "{{empid}}": detaillist[3] if len(detaillist) > 3 else ""
    }

    placeholders2 = {
        "{{research}}": str(research),
        "{{selfm}}": str(selfm),
        "{{mentor}}": str(mentor),
        "{{academics}}": str(academics),
    }

    # also push r/p/s counters into placeholders2
    for i in range(1, 14):
        placeholders2[f"{{{{r{i}_1}}}}"] = r_counters.get(f"r{i}_1", 0)
    for i in range(1, 8):
        placeholders2[f"{{{{p{i}_1}}}}"] = p_counters.get(f"p{i}_1", 0)
    for i in range(1, 6):
        placeholders2[f"{{{{s{i}_1}}}}"] = s_counters.get(f"s{i}_1", 0)


    # Replace placeholders in main doc (paragraph-level)
    try:
        replaced_any = False
        for paragraph in doc.paragraphs:
            text = paragraph.text
            for placeholder, value in placeholders.items():
                if placeholder in text:
                    print(f"Replacing {placeholder} with {value} in paragraph: {text}")
                    text = text.replace(placeholder, str(value))
                    replaced_any = True
            paragraph.text = text
        if not replaced_any:
            print("No placeholders replaced in main doc paragraphs.")
    except Exception as e:
        print("Error replacing placeholders in main doc paragraphs:", e)

    # Also replace inside tables of main doc
    try:
        replaced_table_any = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        for placeholder, value in placeholders.items():
                            if placeholder in text:
                                print(f"Replacing {placeholder} with {value} in table cell: {text}")
                                text = text.replace(placeholder, str(value))
                                replaced_table_any = True
                        paragraph.text = text
        if not replaced_table_any:
            print("No placeholders replaced in main doc tables.")
    except Exception as e:
        print("Error replacing placeholders in main doc tables:", e)

    # Replace placeholders in corrective doc
    try:
        for table in fdoc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = paragraph.text
                        for placeholder, value in placeholders2.items():
                            if placeholder in full_text:
                                full_text = full_text.replace(placeholder, str(value))
                        paragraph.text = full_text
    except Exception as e:
        print("Error replacing placeholders in corrective doc:", e)

    # Additional corrective doc numeric table filling (example logic)
    try:
        if len(fdoc.tables) > 2:
            lasttable = fdoc.tables[2]
            score = [academics, research, selfm, mentor, hod]
            assispro = [0.3, 0.3, 0.15, 0.15, 0.1]
            assospro = [0.2, 0.4, 0.15, 0.15, 0.1]
            prof = [0.1, 0.4, 0.2, 0.2, 0.1]
            tot = 0.0
            designation = detaillist[1] if len(detaillist) > 1 else ""
            # rows 2 to 5 (index 2..5), fill cols 1..5 (index 1..5)
            for row_idx, row in zip(range(2, 6), lasttable.rows[2:6]):
                for cell_idx, cell in zip(range(1, 6), row.cells[1:6]):
                    if row_idx == 2:
                        cell.text = str(score[cell_idx - 1])
                    elif row_idx == 3:
                        if designation == "Professor":
                            cell.text = str(prof[cell_idx - 1])
                        elif designation == "Associate Professor":
                            cell.text = str(assospro[cell_idx - 1])
                        elif designation == "Assistant Professor":
                            cell.text = str(assispro[cell_idx - 1])
                        else:
                            cell.text = "0"
                    elif row_idx == 4:
                        if designation == "Professor":
                            weight = prof[cell_idx - 1]
                        elif designation == "Associate Professor":
                            weight = assospro[cell_idx - 1]
                        elif designation == "Assistant Professor":
                            weight = assispro[cell_idx - 1]
                        else:
                            weight = 0
                        weighted_score = float(score[cell_idx - 1]) * float(weight)
                        cell.text = str(weighted_score)
                        try:
                            tot += weighted_score
                        except Exception:
                            pass
                    else:
                        pass
            # write total to 5th row last cell
            try:
                lasttable.rows[4].cells[-1].text = str(tot)
            except Exception:
                pass
    except Exception as e:
        print("Error filling corrective doc numeric table:", e)

    # Save final docs
    try:
        doc.save("filled_template.docx")
        fdoc.save("appfilled_template.docx")
        # Also save a debug copy
        fdoc.save("debug_filled_template.docx")
        print("Documents saved: filled_template.docx and appfilled_template.docx")
    except Exception as e:
        print("Error saving documents:", e)

    # History append is handled in the upload route


@app.route('/<path:path>')
def static_proxy(path):
    """Serve static files from React build folder."""
    file_path = os.path.join(app.static_folder, path)
    if os.path.exists(file_path):
        return send_from_directory(app.static_folder, path)
    else:
        # Fallback to index.html for client-side routing
        return send_from_directory(app.static_folder, "index.html")

if __name__ == "__main__":
    app.run(debug=True)